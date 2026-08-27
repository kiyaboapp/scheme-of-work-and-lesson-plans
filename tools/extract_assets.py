"""One-off asset inspection helper. Dumps text from the sample docx files and
the first pages of the syllabus / textbook PDFs so the curated data files can be
authored by hand. Not part of the generator package."""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
OUT = ROOT / "tools" / "_extracted"
OUT.mkdir(parents=True, exist_ok=True)


def dump_docx(path: Path, out_name: str) -> None:
    import docx

    doc = docx.Document(str(path))
    lines: list[str] = []
    lines.append(f"=== {path.name} ===")
    lines.append("--- PARAGRAPHS ---")
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(f"[{p.style.name}] {p.text.strip()}")
    lines.append(f"--- TABLES ({len(doc.tables)}) ---")
    for ti, table in enumerate(doc.tables):
        lines.append(f"## table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
            lines.append(f"  r{ri}: " + " || ".join(cells))
    (OUT / out_name).write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {out_name} ({len(lines)} lines)")


def dump_pdf(path: Path, out_name: str, first: int, last: int | None = None) -> None:
    import pdfplumber

    lines: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        total = len(pdf.pages)
        lines.append(f"=== {path.name} :: {total} pages ===")
        end = total if last is None else min(last, total)
        for i in range(first, end):
            text = pdf.pages[i].extract_text() or ""
            lines.append(f"\n----- page {i + 1} -----")
            lines.append(text)
    (OUT / out_name).write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {out_name}")


if __name__ == "__main__":
    what = sys.argv[1] if len(sys.argv) > 1 else "all"
    if what in ("all", "docx"):
        dump_docx(ASSETS / "sample" / "schemeOfWork" / "SCHEME-MATH F1 2026 - W.docx", "sample_scheme.txt")
        dump_docx(ASSETS / "sample" / "lessonPlan" / "MATHEMATICS LESSON FI (1).docx", "sample_lesson.txt")
    if what in ("all", "syllabus"):
        first = int(sys.argv[2]) if len(sys.argv) > 2 else 0
        last = int(sys.argv[3]) if len(sys.argv) > 3 else 30
        dump_pdf(ASSETS / "syllabus" / "MATHEMATICS SYLLABUS - O Level Final.pdf", f"syllabus_{first}_{last}.txt", first, last)
    if what in ("all", "textbook"):
        first = int(sys.argv[2]) if len(sys.argv) > 2 else 0
        last = int(sys.argv[3]) if len(sys.argv) > 3 else 20
        dump_pdf(ASSETS / "textbook" / "MATHEMATICS F1 New - WazaElimu.com.pdf", f"textbook_{first}_{last}.txt", first, last)
