"""Tests for .docx generation endpoints."""

from io import BytesIO

from docx import Document


def _setup_and_allocate(client):
    """Create syllabus, calendar, and run allocation (same pattern as test_scheme.py)."""
    syllabus_data = {
        "subject": "Basic Mathematics",
        "form": "Form One",
        "topics": [
            {
                "order": 1,
                "topic_id_code": "1.0",
                "title": "Demonstrate mastery of mathematical language",
                "sub_topics": [
                    {
                        "order": 1,
                        "sub_topic_id_code": "1.1",
                        "title": "Use numerical skills in different contexts",
                        "planned_periods": 5,
                        "objectives": [
                            {
                                "order": 1,
                                "text": "Explain the basic concepts of Mathematics",
                            },
                            {
                                "order": 2,
                                "text": "Explain the concept of rational, irrational, and real numbers",
                            },
                        ],
                    },
                    {
                        "order": 2,
                        "sub_topic_id_code": "1.2",
                        "title": "Use ratios and proportions in daily life",
                        "planned_periods": 5,
                        "objectives": [
                            {
                                "order": 1,
                                "text": "Explain the concept of ratios and proportions",
                            },
                            {
                                "order": 2,
                                "text": "Solve ratio and proportion problems",
                            },
                        ],
                    },
                ],
            },
        ],
    }

    calendar_data = {
        "academic_year": "2026",
        "terms": [
            {
                "order": 1,
                "term_id_code": "T1",
                "title": "Term 1",
                "weeks": [
                    {
                        "week_number": 1,
                        "start_date": "2026-03-02",
                        "end_date": "2026-03-06",
                        "classification": "teaching",
                        "period_budget": 5,
                        "label": "March Week 1",
                    },
                    {
                        "week_number": 2,
                        "start_date": "2026-03-09",
                        "end_date": "2026-03-13",
                        "classification": "teaching",
                        "period_budget": 5,
                        "label": "March Week 2",
                    },
                ],
            },
        ],
    }

    syl_resp = client.post("/api/syllabus", json=syllabus_data)
    assert syl_resp.status_code == 201
    syllabus_id = syl_resp.json()["id"]

    cal_resp = client.post("/api/calendar", json=calendar_data)
    assert cal_resp.status_code == 201
    calendar_id = cal_resp.json()["id"]

    alloc_resp = client.post(
        "/api/allocate",
        json={"syllabus_id": syllabus_id, "calendar_id": calendar_id},
    )
    assert alloc_resp.status_code == 200

    return syllabus_id, calendar_id


def test_scheme_docx_returns_200_with_correct_content_type(client):
    """GET /api/scheme/{sid}/{cid}/docx returns 200 with correct content-type."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/scheme/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200
    assert (
        resp.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_scheme_docx_content_is_valid_docx(client):
    """The scheme .docx response content is a valid docx file."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/scheme/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200

    # Should be loadable by python-docx
    doc = Document(BytesIO(resp.content))
    assert doc is not None


def test_scheme_docx_has_12_column_table(client):
    """The scheme .docx contains a table with 12 columns."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/scheme/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200

    doc = Document(BytesIO(resp.content))
    # Should have at least one table
    assert len(doc.tables) >= 1

    # First table should have 12 columns
    table = doc.tables[0]
    assert len(table.columns) == 12


def test_scheme_docx_has_correct_headers(client):
    """The scheme .docx has the correct header paragraphs."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/scheme/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200

    doc = Document(BytesIO(resp.content))
    texts = [p.text for p in doc.paragraphs]

    assert "PRESIDENT'S OFFICE" in texts
    assert "REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT" in texts
    assert "SCHEME OF WORK - 2026" in texts


def test_scheme_docx_table_header_row(client):
    """The first row of the scheme table contains expected column headers."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/scheme/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200

    doc = Document(BytesIO(resp.content))
    table = doc.tables[0]
    header_texts = [cell.text for cell in table.rows[0].cells]

    assert "MAIN COMPETENCES" in header_texts
    assert "SPECIFIC COMPETENCES" in header_texts
    assert "LEARNING ACTIVITIES" in header_texts
    assert "PERIODS" in header_texts
    assert "REFERENCES" in header_texts


def test_scheme_docx_has_content_disposition(client):
    """The scheme .docx response has a Content-Disposition header."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/scheme/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200
    assert "content-disposition" in resp.headers
    assert ".docx" in resp.headers["content-disposition"]


def test_lesson_plan_docx_returns_200_with_correct_content_type(client):
    """GET /api/lesson-plan/{sid}/{cid}/docx returns 200 with correct content-type."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/lesson-plan/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200
    assert (
        resp.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_lesson_plan_docx_content_is_valid_docx(client):
    """The lesson plan .docx response content is a valid docx file."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/lesson-plan/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200

    doc = Document(BytesIO(resp.content))
    assert doc is not None


def test_lesson_plan_docx_contains_expected_text(client):
    """The lesson plan .docx contains 'TEACHER'S LESSON PLAN' text."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/lesson-plan/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200

    doc = Document(BytesIO(resp.content))
    all_text = "\n".join([p.text for p in doc.paragraphs])

    assert "TEACHER'S LESSON PLAN" in all_text


def test_lesson_plan_docx_has_stages_table(client):
    """The lesson plan .docx contains a stages table with expected rows."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/lesson-plan/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200

    doc = Document(BytesIO(resp.content))

    # Look for a table with "Introduction" in it (the stages table)
    found_stages_table = False
    for table in doc.tables:
        all_cell_text = " ".join(
            cell.text for row in table.rows for cell in row.cells
        )
        if "Introduction" in all_cell_text and "Competence Development" in all_cell_text:
            found_stages_table = True
            # The stages table has 5 rows (header + 4 stages)
            assert len(table.rows) == 5
            # And 5 columns
            assert len(table.columns) == 5
            break

    assert found_stages_table, "No stages table found in lesson plan docx"


def test_lesson_plan_docx_has_reference_citation(client):
    """The lesson plan .docx contains the TIE reference citation."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/lesson-plan/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200

    doc = Document(BytesIO(resp.content))
    all_text = "\n".join([p.text for p in doc.paragraphs])

    assert "Tanzania Institute of Education" in all_text
    assert "2023" in all_text


def test_lesson_plan_docx_has_content_disposition(client):
    """The lesson plan .docx response has a Content-Disposition header."""
    syllabus_id, calendar_id = _setup_and_allocate(client)

    resp = client.get(f"/api/lesson-plan/{syllabus_id}/{calendar_id}/docx")
    assert resp.status_code == 200
    assert "content-disposition" in resp.headers
    assert ".docx" in resp.headers["content-disposition"]


def test_scheme_docx_404_for_missing_syllabus(client):
    """GET /api/scheme/999/999/docx returns 404 for non-existent data."""
    resp = client.get("/api/scheme/999/999/docx")
    assert resp.status_code == 404


def test_lesson_plan_docx_404_for_missing_syllabus(client):
    """GET /api/lesson-plan/999/999/docx returns 404 for non-existent data."""
    resp = client.get("/api/lesson-plan/999/999/docx")
    assert resp.status_code == 404
