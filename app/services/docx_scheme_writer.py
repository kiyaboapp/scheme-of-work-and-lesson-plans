"""DOCX writer for Scheme of Work documents."""

from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.schemas import SchemeOfWorkResponse


# Column headers matching the sample scheme format
SCHEME_COLUMNS = [
    "MAIN COMPETENCES",
    "SPECIFIC COMPETENCES",
    "LEARNING ACTIVITIES",
    "SPECIFIC ACTIVITIES",
    "MONTH",
    "WEEK",
    "PERIODS",
    "TEACHING AND LEARNING METHODS",
    "TEACHING AND LEARNING RESOURCES",
    "ASSESSMENT TOOLS",
    "REFERENCES",
    "REMARK",
]

# Map week_number to ordinal text
ORDINAL_MAP = {
    1: "1st",
    2: "2nd",
    3: "3rd",
    4: "4th",
    5: "5th",
    6: "6th",
    7: "7th",
    8: "8th",
    9: "9th",
    10: "10th",
}


def _week_to_ordinal(week_number: int) -> str:
    """Convert a week number to ordinal text."""
    return ORDINAL_MAP.get(week_number, f"{week_number}th")


def _derive_month(week_label: str | None) -> str:
    """Derive month name from week_label."""
    if not week_label:
        return ""
    # Week labels may contain month names or date info
    label_upper = week_label.upper()
    months = [
        "JANUARY", "FEBRUARY", "MARCH", "APRIL", "MAY", "JUNE",
        "JULY", "AUGUST", "SEPTEMBER", "OCTOBER", "NOVEMBER", "DECEMBER",
    ]
    for month in months:
        if month in label_upper:
            return month
    # Try abbreviated versions
    abbrevs = {
        "JAN": "JANUARY", "FEB": "FEBRUARY", "MAR": "MARCH",
        "APR": "APRIL", "MAY": "MAY", "JUN": "JUNE",
        "JUL": "JULY", "AUG": "AUGUST", "SEP": "SEPTEMBER",
        "SEPT": "SEPTEMBER", "OCT": "OCTOBER", "NOV": "NOVEMBER",
        "DEC": "DECEMBER",
    }
    for abbrev, full in abbrevs.items():
        if abbrev in label_upper:
            return full
    return week_label.upper() if week_label else ""


def _add_full_width_row(table, text: str):
    """Add a full-width row with the same text in all 12 columns (mimics merged cells)."""
    row = table.add_row()
    for cell in row.cells:
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)


def generate_scheme_docx(scheme: SchemeOfWorkResponse) -> BytesIO:
    """
    Generate a .docx file for the given Scheme of Work response.

    Returns a BytesIO object containing the document.
    """
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # Header paragraphs matching the sample
    p1 = doc.add_paragraph("PRESIDENT'S OFFICE")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p1.runs:
        run.bold = True

    p2 = doc.add_paragraph("REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.bold = True

    p3 = doc.add_paragraph(f"SCHEME OF WORK - {scheme.academic_year}")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p3.runs:
        run.bold = True

    # Teacher/School/Subject info
    doc.add_paragraph("TEACHER'S NAME: .........")
    doc.add_paragraph("SCHOOL'S NAME: .........")
    doc.add_paragraph(f"SUBJECT: {scheme.subject.upper()}")
    doc.add_paragraph(f"CLASS: {scheme.form.upper()}")
    doc.add_paragraph("TERM: 1st & 2nd")
    doc.add_paragraph(f"YEAR: {scheme.academic_year}")

    # Create the main table
    table = doc.add_table(rows=1, cols=12)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_row = table.rows[0]
    for idx, col_name in enumerate(SCHEME_COLUMNS):
        cell = header_row.cells[idx]
        cell.text = col_name
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    # Track which term we are in to insert break/assessment rows
    current_term = None
    prev_term = None
    term_teaching_count = 0
    midterm_inserted = set()
    terminal_inserted = set()

    # Group entries by term to determine midterm and terminal positions
    term_entries: dict[str, list] = {}
    for entry in scheme.entries:
        if entry.term not in term_entries:
            term_entries[entry.term] = []
        term_entries[entry.term].append(entry)

    # Process entries and add special rows
    for i, entry in enumerate(scheme.entries):
        current_term = entry.term

        # Insert midterm assessment/break at halfway point for each term
        if current_term != prev_term:
            term_teaching_count = 0

            # If switching terms, add terminal assessment + break for previous term
            if prev_term and prev_term not in terminal_inserted:
                terminal_inserted.add(prev_term)
                _add_full_width_row(table, "TERMINAL ASSESSMENT")
                _add_full_width_row(table, "TERM BREAK")

        term_teaching_count += 1
        term_total = len(term_entries.get(current_term, []))

        # Insert midterm at roughly the halfway point
        if (
            current_term not in midterm_inserted
            and term_total > 2
            and term_teaching_count == (term_total // 2) + 1
        ):
            midterm_inserted.add(current_term)
            _add_full_width_row(table, "MID-TERM ASSESSMENT")
            _add_full_width_row(table, "MID-TERM BREAK")

        # Data row
        row = table.add_row()
        cells = row.cells

        # MAIN COMPETENCES = topic_title
        cells[0].text = entry.topic_title

        # SPECIFIC COMPETENCES = sub_topic_title
        cells[1].text = entry.sub_topic_title

        # LEARNING ACTIVITIES = first objective
        if entry.objectives:
            cells[2].text = entry.objectives[0]
        else:
            cells[2].text = entry.sub_topic_title

        # SPECIFIC ACTIVITIES = teacher-facing derived from objectives
        if entry.objectives:
            activities = []
            for obj in entry.objectives:
                activities.append(f"To guide students to {obj.lower()}")
            cells[3].text = "\n".join(activities)
        else:
            cells[3].text = f"To guide students to {entry.sub_topic_title.lower()}"

        # MONTH
        cells[4].text = _derive_month(entry.week_label)

        # WEEK = ordinal
        cells[5].text = _week_to_ordinal(entry.week_number)

        # PERIODS
        cells[6].text = str(entry.periods)

        # TEACHING AND LEARNING METHODS
        cells[7].text = entry.teaching_methods or "Presentations, Problem solving, Practical work"

        # TEACHING AND LEARNING RESOURCES
        cells[8].text = entry.teaching_resources or ""

        # ASSESSMENT TOOLS
        cells[9].text = "Oral Questions Exercises Quizzes Home works"

        # REFERENCES
        cells[10].text = entry.references or ""

        # REMARK
        cells[11].text = ""

        # Set font size for data cells
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

        prev_term = current_term

    # Add final terminal assessment and break for the last term
    if prev_term and prev_term not in terminal_inserted:
        _add_full_width_row(table, "REVISION AND ANNUAL ASSESSMENT")
        _add_full_width_row(table, f"END OF YEAR BREAK")

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
