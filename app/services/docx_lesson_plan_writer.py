"""DOCX writer for Lesson Plan documents."""

from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.models import AllocationAssignment, LessonPlan


# Reference citation matching the sample
REFERENCE_CITATION = (
    "Tanzania Institute of Education. (2023). Basic Mathematics for secondary "
    "schools student's book, Form One. Tanzania Institute of Education."
)


def _get_objective_for_period(assignment: AllocationAssignment, period_number: int) -> str:
    """Get the objective text for a given period within an assignment."""
    objectives = sorted(assignment.sub_topic.objectives, key=lambda o: o.order)
    if not objectives:
        return assignment.sub_topic.title

    # Map period to objective: repeat objectives across periods
    idx = (period_number - assignment.first_period) % len(objectives)
    return objectives[idx].text


def _add_lesson_plan_section(
    doc: Document,
    assignment: AllocationAssignment,
    lesson_plan: LessonPlan,
    subject: str,
    form: str,
):
    """Add a single lesson plan section to the document."""
    sub_topic = assignment.sub_topic
    topic = sub_topic.topic

    # Main competence and specific competence
    main_competence = topic.title
    specific_competence = sub_topic.title

    # Get objective for this period
    objective_text = _get_objective_for_period(assignment, lesson_plan.period_number)

    # Title
    title_p = doc.add_paragraph("TEACHER'S LESSON PLAN")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.bold = True
        run.font.size = Pt(12)

    # Header fields
    doc.add_paragraph(
        "Name of School: .........      "
        "Teacher's Name: ........."
    )
    doc.add_paragraph(
        f"Form: {form}                                        "
        f"Subject: {subject}"
    )

    date_str = str(lesson_plan.date) if lesson_plan.date else "........."
    doc.add_paragraph(
        f"Time: .........                    "
        f"Date: {date_str}"
    )

    # Competence table (2 rows x 2 cols)
    comp_table = doc.add_table(rows=2, cols=2)
    comp_table.style = "Table Grid"
    comp_table.rows[0].cells[0].text = "Main Competence:"
    comp_table.rows[0].cells[1].text = main_competence
    comp_table.rows[1].cells[0].text = "Specific Competence:"
    comp_table.rows[1].cells[1].text = specific_competence

    doc.add_paragraph("")  # Spacer

    # Main Activity
    main_activity_p = doc.add_paragraph(
        f"Main Activity: Within 1 period students should be able to "
        f"{specific_competence.lower()}"
    )
    for run in main_activity_p.runs:
        run.font.size = Pt(10)

    # Specific Activity
    specific_activity_p = doc.add_paragraph(
        f"Specific Activity: Within 40 minutes, students should be to;"
    )
    for run in specific_activity_p.runs:
        run.font.size = Pt(10)

    # Objective text
    obj_p = doc.add_paragraph(objective_text)
    for run in obj_p.runs:
        run.font.size = Pt(10)

    # Teaching and Learning Resources
    resources_text = lesson_plan.teaching_resources or "Charts, calculators, real objects"
    doc.add_paragraph(f"Teaching and Learning Resources:")
    doc.add_paragraph(resources_text)

    # References
    doc.add_paragraph(f"References: {REFERENCE_CITATION}")

    # Teaching and Learning Process header
    process_p = doc.add_paragraph("Teaching and Learning Process")
    process_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in process_p.runs:
        run.bold = True

    # Stages table (5 rows x 5 cols: header + 4 stage rows)
    stages_table = doc.add_table(rows=5, cols=5)
    stages_table.style = "Table Grid"

    # Header row
    headers = ["Stages", "Time (Minutes)", "Teaching Activities", "Learning Activities", "Assessment Criteria"]
    for idx, header_text in enumerate(headers):
        cell = stages_table.rows[0].cells[idx]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Introduction row
    stages_table.rows[1].cells[0].text = "Introduction"
    stages_table.rows[1].cells[1].text = "05"
    stages_table.rows[1].cells[2].text = (
        "Displaying pictures on manila sheet. "
        "Asking students some questions about today's lesson."
    )
    stages_table.rows[1].cells[3].text = "Observe the image and respond to the questions asked."
    stages_table.rows[1].cells[4].text = "Questions about today's lesson are answered."

    # Competence Development row
    stages_table.rows[2].cells[0].text = "Competence Development"
    stages_table.rows[2].cells[1].text = "20"
    stages_table.rows[2].cells[2].text = (
        f"Provide students with short guiding questions and ask them in groups to "
        f"{objective_text.lower()}. "
        "Display a video or pictures with different activities of a lesson. "
        "Then, ask them in pairs to identify activities the video or pictures drawn."
    )
    stages_table.rows[2].cells[3].text = (
        "Discuss and share what they have learnt in the lesson. "
        "Watch and identify different activities of a lesson from the video."
    )
    stages_table.rows[2].cells[4].text = (
        "Everything taught today is clearly explained. "
        "Different activities of a lesson are identified."
    )

    # Design row
    stages_table.rows[3].cells[0].text = "Design"
    stages_table.rows[3].cells[1].text = "10"
    stages_table.rows[3].cells[2].text = (
        "Ask students in groups to name and explain different things "
        "they have learnt when participating in discussion."
    )
    stages_table.rows[3].cells[3].text = (
        "Name and explain different things they have learnt "
        "when participating in discussion."
    )
    stages_table.rows[3].cells[4].text = "Different things learnt today are identified."

    # Realisation row
    stages_table.rows[4].cells[0].text = "Realisation"
    stages_table.rows[4].cells[1].text = "05"
    stages_table.rows[4].cells[2].text = (
        "Ask each student to show and give examples of what "
        "she or he has learnt in the lesson today."
    )
    stages_table.rows[4].cells[3].text = (
        "Show and give examples of what she or he has learnt "
        "in the lesson today."
    )
    stages_table.rows[4].cells[4].text = "Examples are given and things learnt are explained."

    # Remarks section
    doc.add_paragraph("")
    remarks_text = lesson_plan.remarks or (
        "The students were able to explain things taught today due to the use of "
        "interactive teaching and learning methods, activities and resources."
    )
    doc.add_paragraph(f"Remarks: {remarks_text}")

    # Page break between lesson plans
    doc.add_page_break()


def generate_lesson_plan_docx(
    assignments: list[AllocationAssignment],
    lesson_plans_by_assignment: dict[int, list[LessonPlan]],
    subject: str,
    form: str,
) -> BytesIO:
    """
    Generate a .docx file containing all lesson plans.

    Args:
        assignments: Ordered list of allocation assignments
        lesson_plans_by_assignment: Dict mapping assignment_id -> list of LessonPlan
        subject: Subject name (e.g. "Basic Mathematics")
        form: Form/class name (e.g. "One")

    Returns a BytesIO object containing the document.
    """
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    for assignment in assignments:
        plans = lesson_plans_by_assignment.get(assignment.id, [])
        for plan in sorted(plans, key=lambda p: p.period_number):
            _add_lesson_plan_section(doc, assignment, plan, subject, form)

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
